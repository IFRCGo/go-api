import datetime
import docx
from typing import Any, List

from dref.models import Dref


def parse_int(s):
    try:
        return int(s)
    except (ValueError, TypeError):
        return None


def parse_date(date):
    formats = ['%d/%m/%Y', '%Y-%m-%d']
    for fmt in formats:
        try:
            return datetime.datetime.strptime(date, fmt)
        except (ValueError, TypeError):
            pass
    return None


def parse_float(s):
    try:
        return float(s)
    except (ValueError, TypeError):
        return None

def get_text_or_null(colitems: List[Any]):
    if colitems:
        return colitems[0].text
    return None


def parse_boolean_or_null(colitems: List[Any]):
    if colitems:
        return parse_boolean(colitems[0].text)
    return None


def parse_string_to_int(string):
    try:
        char_to_check = ','
        if string and char_to_check in string:
            new_strings = string.split(',')
            concat_string = new_strings[0] + new_strings[1]
            return int(concat_string)
        return int(string)
    except (ValueError, TypeError):
        return None


def parse_boolean(string):
    if string and string == 'Yes':
        return True
    elif string and string == 'No':
        return False
    return None


def get_table_data(doc):
    document = docx.Document(doc)
    tables = []
    for table in document.tables:
        rowdata = []
        for _, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cells.append([x.text for x in cell._tc.xpath('.//w:t')])
            rowdata.append(cells)
        tables.append(rowdata)
    return tables

def get_paragraphs_data(doc):
    document = docx.Document(doc)
    return [[y.text for y in x._element.xpath('.//w:t')] for x in document.paragraphs]


def parse_disaster_category(disaster_category):
    if disaster_category == 'Yellow':
        return Dref.DisasterCategory.YELLOW
    elif disaster_category == 'Orange':
        return Dref.DisasterCategory.ORANGE
    elif disaster_category == 'Red':
        return Dref.DisasterCategory.RED
    return None
