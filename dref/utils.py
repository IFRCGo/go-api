import logging
import docx

from typing import List, Any

from rest_framework import serializers

from dref.models import (
    Dref,
    PlannedIntervention,
    IdentifiedNeed,
    NationalSocietyAction,
    PlannedInterventionIndicators,
    RiskSecurity,
)
from api.models import (
    DisasterType,
    Country,
)

from .common_utils import (
    parse_int,
    parse_string_to_int,
    parse_date,
    parse_float,
    parse_boolean,
    get_table_data,
    get_paragraphs_data,
    parse_disaster_category,
)


logger = logging.getLogger(__name__)


def parse_national_society_title(title):
    title_dict = {
        'National Society Readiness': NationalSocietyAction.Title.NATIONAL_SOCIETY_READINESS,
        'Assessments': NationalSocietyAction.Title.ASSESSMENT,
        'Coordination': NationalSocietyAction.Title.COORDINATION,
        'Resource Mobilization': NationalSocietyAction.Title.RESOURCE_MOBILIZATION,
        'Activation of contingency plans': NationalSocietyAction.Title.ACTIVATION_OF_CONTINGENCY_PLANS,
        'EOC': NationalSocietyAction.Title.NATIONAL_SOCIETY_EOC,
        'Shelter, Housing and Settlements': NationalSocietyAction.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Livelihoods and Basic Needs': NationalSocietyAction.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health': NationalSocietyAction.Title.HEALTH,
        'Water Sanitation and Hygiene': NationalSocietyAction.Title.WATER_SANITATION_AND_HYGIENE,
        'Protection, Gender and Inclusion': NationalSocietyAction.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': NationalSocietyAction.Title.EDUCATION,
        'Migration': NationalSocietyAction.Title.MIGRATION,
        'Risk Reduction, Climate Adaptation and Recovery': NationalSocietyAction.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Community Engagement and Accountability': NationalSocietyAction.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        'Environment Sustainability': NationalSocietyAction.Title.ENVIRONMENT_SUSTAINABILITY,
        'Multi-purpose Cash': NationalSocietyAction.Title.MULTI_PURPOSE_CASH,
        'Other': NationalSocietyAction.Title.OTHER,

    }
    return title_dict.get(title)


def parse_identified_need_title(title):
    title_dict = {
        'Livelihoods': IdentifiedNeed.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health & Care': IdentifiedNeed.Title.HEALTH,
        'Water, Sanitation and Hygiene (WASH)': IdentifiedNeed.Title.WATER_SANITATION_AND_HYGIENE,
        'Risk Reduction, Climate adaptation and Recovery': IdentifiedNeed.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Shelter, Housing and Settlements': IdentifiedNeed.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Multi-purpose Cash': IdentifiedNeed.Title.MULTI_PURPOSE_CASH_GRANTS,
        'Protection, Gender, and Inclusion (PGI)': IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': IdentifiedNeed.Title.EDUCATION,
        'Migration': IdentifiedNeed.Title.MIGRATION,
        'Environmental Sustainability': IdentifiedNeed.Title.ENVIRONMENT_SUSTAINABILITY,
        'Shelter Cluster Coordination': IdentifiedNeed.Title.SHELTER_CLUSTER_COORDINATION
    }
    return title_dict.get(title)


def parse_planned_intervention_title(title):
    title_dict = {
        'Shelter, Housing and Settlements': PlannedIntervention.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Livelihoods And Basic Needs': PlannedIntervention.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health': PlannedIntervention.Title.HEALTH,
        'Water, Sanitation And Hygiene': PlannedIntervention.Title.WATER_SANITATION_AND_HYGIENE,
        'Protection, Gender And Inclusion': PlannedIntervention.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': PlannedIntervention.Title.EDUCATION,
        'Migration': PlannedIntervention.Title.MIGRATION,
        'Risk Reduction, Climate Adaptation And Recovery': PlannedIntervention.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Environmental Sustainability': PlannedIntervention.Title.ENVIRONMENTAL_SUSTAINABILITY,
        'Secretariat Services': PlannedIntervention.Title.SECRETARIAT_SERVICES,
        'National Society Strengthening': PlannedIntervention.Title.NATIONAL_SOCIETY_STRENGTHENING,
        'Multi-purpose Cash': PlannedIntervention.Title.MULTI_PURPOSE_CASH,
        'Community Engagement and Accountability': PlannedIntervention.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,


    }
    return title_dict.get(title)


def parse_disaster_type(disaster_type):
    return DisasterType.objects.filter(name__icontains=disaster_type).first()


def parse_type_of_onset(type):
    if type == 'Slow':
        type = Dref.OnsetType.SLOW
    elif type == 'Sudden':
        type = Dref.OnsetType.SUDDEN
    elif type == 'Imminent':
        type = Dref.OnsetType.Imminent
    return type


def extract_file(doc, created_by):
    document = docx.Document(doc)
    tables = get_table_data(doc)
    paragraphs = get_paragraphs_data(doc)

    data = {}

    # Get item form an array
    def _t(items: List[Any], index=0):
        if items and len(items) > index:
            return items[index]
        return None

    def get_table_cells(table_idx):
        table = tables[table_idx]
        def f(r, c, i=0, as_list=False):
            if as_list:
                return table[r][c]
            return _t(table[r][c], i)
        return f

    # NOTE: Second Paragraph for Country and Region and Dref Title
    title_paragraph = paragraphs[1]
    if len(title_paragraph) > 0:
        data['title'] = title_paragraph[0]
    else:
        # FIXME: raise validation error
        data['title'] = None
        # raise serializers.ValidationError('Title is required to create dref')

    table = document.tables[0]
    cells = get_table_cells(0)
    try:
        data['appeal_code'] = cells(1, 0)
    except(IndexError, ValueError):
        pass
    try:
        data['amount_requested'] = parse_string_to_int(cells(1, 1, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['disaster_category'] = parse_disaster_category(cells(1, 2))
    except(IndexError, ValueError):
        pass
    try:
        data['disaster_type'] = parse_disaster_type(cells(1, 4))
    except(IndexError, ValueError):
        pass
    try:
        data['glide_code'] = cells(3, 0)
    except(IndexError, ValueError):
        pass
    try:
        data['num_affected'] = parse_string_to_int(cells(3, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['num_assisted'] = parse_string_to_int(cells(3, 2))
    except(IndexError, ValueError):
        pass
    try:
        data['type_of_onset'] = parse_type_of_onset(cells(5, 0))
    except(IndexError, ValueError):
        pass
    try:
        data['date_of_approval'] = parse_date(cells(5, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['end_date'] = parse_date(cells(5, 2))
    except(IndexError, ValueError):
        pass
    try:
        data['operation_timeframe'] = parse_int(cells(5, 3))
    except(IndexError, ValueError):
        pass
    country_name = cells(6, 0, 1)
    try:
        data['country'] = Country.objects.get(name_en__icontains=country_name)
    except IndexError:
        raise serializers.ValidationError('Country is required')
    # FIXME: raise validation error if no country
    except Country.DoesNotExist:
        data['country'] = None

    try:
        description = paragraphs[7]
        data['event_description'] = ''.join(description) if description else None
    except (IndexError, ValueError):
        pass
    try:
        event_scope = paragraphs[9]
        data['event_scope'] = ''.join(event_scope) if event_scope else None
    except(IndexError, ValueError):
        pass
    # Previous Operation
    cells = get_table_cells(1)
    try:
        data['affect_same_area'] = parse_boolean(cells(0, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['affect_same_population'] = parse_boolean(cells(1, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['ns_respond'] = parse_boolean(cells(2, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['ns_request_fund'] = parse_boolean(cells(3, 1))
    except(IndexError, ValueError):
        pass
    try:
        data['ns_request_text'] = cells(4, 1)
        if cells(4, 1) == 'Yes':
            data['dref_recurrent_text'] = cells(5, 1)
    except(IndexError, ValueError):
        pass
    try:
        if cells(0, 1) == 'Yes' and cells(1, 1) == 'Yes' and cells(2, 1) == 'Yes' and cells(3, 1) == 'Yes' and cells(4, 1) == 'Yes':
            recurrent_text = cells(7, 1, as_list=True) or []
            data['dref_recurrent_text'] = ''.join(recurrent_text) if recurrent_text else None
    except(IndexError, ValueError):
        pass
    try:
        data['lessons_learned'] = cells(8, 1)
    except IndexError:
        pass
    ## Movement Parameters
    cells = get_table_cells(3)
    try:
        ifrc_desc = cells(0, 1, as_list=True) or []
        data['ifrc'] = ''.join(ifrc_desc) if ifrc_desc else None
        partner_national_society_desc = cells(1, 1, as_list=True) or []
        data['partner_national_society'] = ''.join(partner_national_society_desc) if partner_national_society_desc else None
        icrc_desc = cells(2, 1, as_list=True) or []
        data['icrc'] = ''.join(icrc_desc) if icrc_desc else None
    except IndexError:
        pass

    # National Socierty Actions
    action_titles = [
        NationalSocietyAction.Title.NATIONAL_SOCIETY_READINESS,
        NationalSocietyAction.Title.ASSESSMENT,
        NationalSocietyAction.Title.COORDINATION,
        NationalSocietyAction.Title.RESOURCE_MOBILIZATION,
        NationalSocietyAction.Title.ACTIVATION_OF_CONTINGENCY_PLANS,
        NationalSocietyAction.Title.NATIONAL_SOCIETY_EOC,
        NationalSocietyAction.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        NationalSocietyAction.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        NationalSocietyAction.Title.HEALTH,
        NationalSocietyAction.Title.WATER_SANITATION_AND_HYGIENE,
        NationalSocietyAction.Title.PROTECTION_GENDER_AND_INCLUSION,
        NationalSocietyAction.Title.EDUCATION,
        NationalSocietyAction.Title.MIGRATION,
        NationalSocietyAction.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        NationalSocietyAction.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        NationalSocietyAction.Title.ENVIRONMENT_SUSTAINABILITY,
        NationalSocietyAction.Title.MULTI_PURPOSE_CASH,
        NationalSocietyAction.Title.OTHER,
    ]

    cells = get_table_cells(2)
    # National Society
    national_society_actions = []
    for i, title in enumerate(action_titles):
        try:
            description = cells(i, 1)
            if description:
                data_new = {
                    'title': title,
                    'description': description,
                }
                national_society_actions.append(data_new)
        except IndexError:
            pass

    # Crete national Society objects db level
    national_societys = []
    for national_data in national_society_actions:
        if national_data['description']:
            national = NationalSocietyAction.objects.create(**national_data)
            national_societys.append(national)
    # Other actors
    table4 = document.tables[4]
    cells = get_table_cells(4)
    try:
        if cells(0, 0):
            data['government_requested_assistance'] = parse_boolean(cells(0, 1))
    except (ValueError, IndexError):
        pass
    try:
        national_authorities = cells(1, 1, as_list=True) or []
        data['national_authorities'] = ''.join(national_authorities) if national_authorities else None
    except (ValueError, IndexError):
        pass

    try:
        un_and_other_actors = cells(1, 2, as_list=True) or []
        data['un_or_other_actor'] = ''.join(un_and_other_actors) if un_and_other_actors else None
    except (IndexError, ValueError):
        pass

    try:
        coordination_mechanism = cells(3, 1, as_list=True) or []
        data['major_coordination_mechanism'] = ''.join(coordination_mechanism) if coordination_mechanism else None
    except (IndexError, ValueError):
        pass
    # NeedsIdentified
    table5 = document.tables[5]
    cells = get_table_cells(5)
    needs_identified = []
    needs_titles = [
        IdentifiedNeed.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        IdentifiedNeed.Title.HEALTH,
        IdentifiedNeed.Title.WATER_SANITATION_AND_HYGIENE,
        IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        IdentifiedNeed.Title.EDUCATION,
        IdentifiedNeed.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        IdentifiedNeed.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        IdentifiedNeed.Title.ENVIRONMENT_SUSTAINABILITY,
        IdentifiedNeed.Title.SHELTER_CLUSTER_COORDINATION,
    ]
    for i, needs_title in enumerate(needs_titles):
        try:
            description = cells(i*2+1, 0)  # Use Every alternate row, so i*2 + 1
            if description:
                data_new = {
                    'title': needs_title,
                    'description': description
                }
                needs_identified.append(data_new)
        except IndexError:
            pass

    needs = []
    for need in needs_identified:
        if need['description']:
            identified = IdentifiedNeed.objects.create(**need)
            needs.append(identified)
    try:
        data['operation_objective'] = ''.join(paragraphs[23] or [])
    except (IndexError, ValueError):
        pass

    # targeting strategy
    try:
        data['response_strategy'] = ''.join(paragraphs[25] or [])
    except (IndexError, ValueError):
        pass
    # Targeting Strategy
    try:
        data['people_assisted'] = ''.join(paragraphs[31] or [])
    except(IndexError, ValueError):
        pass
    try:
        data['selection_criteria'] = ''.join(paragraphs[33] or [])
    except(IndexError, ValueError):
        pass

    # Targeting Population
    cells = get_table_cells(6)
    categories = ['women', 'girls', 'men', 'boys', 'total_targeted_population']
    for i, category in enumerate(categories):
        try:
            data[category] = parse_string_to_int(cells(i, 1))
        except(IndexError, ValueError):
            pass

    try:
        data['people_per_local'] = parse_float(cells(1, 2))
    except(IndexError, ValueError):
        pass
    try:
        data['people_per_urban'] = parse_float(cells(1, 3))
    except(IndexError, ValueError):
        pass
    try:
        data['disability_people_per'] = parse_float(cells(3, 2))
    except(IndexError, ValueError):
        pass

    # Risk And Security Considerations
    cells = get_table_cells(7)
    mitigation_actions_rows = range(2, 5)
    mitigation_actions = []
    for i in mitigation_actions_rows:
        try:
            data_row_one = {
                'risk': cells(i, 0),
                'mitigation': cells(i, 1),
            }
            mitigation_actions.append(data_row_one)
        except(IndexError, ValueError):
            pass

    try:
        data['risk_security_concern'] = cells(6, 0)
    except IndexError:
        pass

    risk_security_list = []
    for mitigation in mitigation_actions:
        if mitigation['risk'] or mitigation['mitigation']:
            risk_security = RiskSecurity.objects.create(**mitigation)
            risk_security_list.append(risk_security)

    # About Support Service
    data['human_resource'] = ''.join(paragraphs[57] or [])
    data['surge_personnel_deployed'] = ''.join(paragraphs[60] or [])
    data['logistic_capacity_of_ns'] = ''.join(paragraphs[62] or [])
    data['pmer'] = ''.join(paragraphs[64] or [])
    data['communication'] = ''.join(paragraphs[66] or [])

    try:
        national_society_contact = paragraphs[71] or []
        data['national_society_contact_title'] = national_society_contact[3] if national_society_contact[3] else None
        data['national_society_contact_email'] = national_society_contact[5] if national_society_contact[5] else None
        data['national_society_contact_phone_number'] = national_society_contact[7] if national_society_contact[7] else None
        data['national_society_contact_name'] = national_society_contact[0] if national_society_contact[0] else None
    except IndexError:
        pass
    try:

        ifrc_appeal_manager = paragraphs[72] or []
        data['ifrc_appeal_manager_title'] = ifrc_appeal_manager[3] if ifrc_appeal_manager[3] else None
        data['ifrc_appeal_manager_email'] = ifrc_appeal_manager[5] if ifrc_appeal_manager[5] else None
        data['ifrc_appeal_manager_phone_number'] = ifrc_appeal_manager[7] if ifrc_appeal_manager[7] else None
        data['ifrc_appeal_manager_name'] = ifrc_appeal_manager[0] if ifrc_appeal_manager[0] else None
    except IndexError:
        pass

    try:
        ifrc_project_manager = paragraphs[73] or []
        data['ifrc_project_manager_title'] = ifrc_project_manager[3] if ifrc_project_manager[3] else None
        data['ifrc_project_manager_email'] = ifrc_project_manager[5] if ifrc_project_manager[5] else None
        data['ifrc_project_manager_phone_number'] = ifrc_project_manager[7] if ifrc_project_manager[7] else None
        data['ifrc_project_manager_name'] = ifrc_project_manager[0] if ifrc_project_manager[0] else None
    except IndexError:
        pass

    try:
        ifrc_emergency = paragraphs[74] or []
        data['ifrc_emergency_title'] = ifrc_emergency[3] if ifrc_emergency[3] else None
        data['ifrc_emergency_email'] = ifrc_emergency[5] if ifrc_emergency[5] else None
        data['ifrc_emergency_phone_number'] = ifrc_emergency[7] if ifrc_emergency[7] else None
        data['ifrc_emergency_name'] = ifrc_emergency[0] if ifrc_emergency[0] else None
    except IndexError:
        pass

    try:
        media = paragraphs[75] or []
        data['media_contact_title'] = media[3] if media[3] else None
        data['media_contact_email'] = media[5] if media[5] else None
        data['media_contact_phone_number'] = media[7] if media[7] else None
        data['media_contact_name'] = media[0] if media[0] else None
    except IndexError:
        pass

    # PlannedIntervention Table
    planned_intervention = []
    for i in range(8, 19):
        table = document.tables[i]
        cells = get_table_cells(i)
        title = cells(0, 1)
        budget = cells(0, 3, 1)
        targated_population = cells(1, 3)

        indicators = []
        for j in range(3, 6):
            try:
                indicators.append({
                    'title': cells(j, 0),
                    'target': cells(j, 2),
                })
            except(IndexError, ValueError):
                pass

        indicators_object_list = []
        for indicator in indicators:
            if indicator['title']:
                try:
                    planned_object = PlannedInterventionIndicators.objects.create(**indicator)
                    indicators_object_list.append(planned_object)
                except ValueError:  # FIXME: this try except should ideally not exist
                    logger.warning(f'Could not insert planned object for indicator {indicator}')
                    pass
        try:
            description_text = table.cell(6, 1)._tc.xpath('.//w:t')
            priority_description = description_text[0].text
            planned_data = {
                'title': parse_planned_intervention_title(title),
                'budget': budget,
                'person_targeted': targated_population,
                'description': priority_description
            }

            planned = PlannedIntervention.objects.create(**planned_data)
            planned.indicators.add(*indicators_object_list)
            planned_intervention.append(planned)
        except(IndexError, ValueError):
            pass

    # Create dref objects
    # map m2m fields
    data['is_published'] = False
    data['national_society'] = Country.objects.filter(name_en__icontains=country_name).first()
    data['created_by'] = created_by
    dref = Dref.objects.create(**data)
    dref.planned_interventions.add(*planned_intervention)
    dref.needs_identified.add(*needs)
    dref.national_society_actions.add(*national_societys)
    dref.risk_security.add(*risk_security_list)
    return dref
