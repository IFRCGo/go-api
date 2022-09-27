import docx
from rest_framework import serializers

from typing import Any, List

from dref.models import (
    Dref,
    PlannedIntervention,
    IdentifiedNeed,
    NationalSocietyAction,
    PlannedInterventionIndicators,
    RiskSecurity
)
from api.models import (
    DisasterType,
    Country,
    District
)

from .common_utils import (
    parse_int,
    parse_date,
    parse_float,
    get_text_or_null,
    parse_string_to_int,
    parse_boolean,
    get_table_data,
    parse_disaster_category,
    get_paragraphs_data,
    parse_contact_information,
)


def parse_national_society_title(title):
    title_dict = {
        'National Society Readiness': NationalSocietyAction.Title.NATIONAL_SOCIETY_READINESS,
        'Assessments': NationalSocietyAction.Title.ASSESSMENT,
        'Coordination': NationalSocietyAction.Title.COORDINATION,
        'Resource Mobilization': NationalSocietyAction.Title.RESOURCE_MOBILIZATION,
        'Activation of contingency plans': NationalSocietyAction.Title.ACTIVATION_OF_CONTINGENCY_PLANS,
        'EOC': NationalSocietyAction.Title.NATIONAL_SOCIETY_EOC,
        'Shelter, Housing and Settlements': NationalSocietyAction.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Livelihoods and Basic Needs': NationalSocietyAction.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health': NationalSocietyAction.Title.HEALTH,
        'Water Sanitation and Hygiene': NationalSocietyAction.Title.WATER_SANITATION_AND_HYGIENE,
        'Protection, Gender and Inclusion': NationalSocietyAction.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': NationalSocietyAction.Title.EDUCATION,
        'Migration': NationalSocietyAction.Title.MIGRATION,
        'Risk Reduction, Climate Adaptation and Recovery': NationalSocietyAction.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Community Engagement and Accountability': NationalSocietyAction.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        'Environment Sustainability': NationalSocietyAction.Title.ENVIRONMENT_SUSTAINABILITY,
        'Other': NationalSocietyAction.Title.OTHER,

    }
    return title_dict.get(title)


def parse_identified_need_title(title):
    title_dict = {
        'Livelihoods': IdentifiedNeed.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health & Care': IdentifiedNeed.Title.HEALTH,
        'Water, Sanitation and Hygiene (WASH)': IdentifiedNeed.Title.WATER_SANITATION_AND_HYGIENE,
        'Risk Reduction, Climate adaptation and Recovery': IdentifiedNeed.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Shelter, Housing and Settlements': IdentifiedNeed.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Multi-purpose Cash': IdentifiedNeed.Title.MULTI_PURPOSE_CASH_GRANTS,
        'Protection, Gender, and Inclusion (PGI)': IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': IdentifiedNeed.Title.EDUCATION,
        'Migration': IdentifiedNeed.Title.MIGRATION,
        'Environmental Sustainability': IdentifiedNeed.Title.ENVIRONMENT_SUSTAINABILITY,
    }
    return title_dict.get(title)


def parse_planned_intervention_title(title):
    title_dict = {
        'Shelter, Housing and Settlements': PlannedIntervention.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        'Livelihoods And Basic Needs': PlannedIntervention.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        'Health': PlannedIntervention.Title.HEALTH,
        'Water, Sanitation And Hygiene': PlannedIntervention.Title.WATER_SANITATION_AND_HYGIENE,
        'Protection, Gender And Inclusion': PlannedIntervention.Title.PROTECTION_GENDER_AND_INCLUSION,
        'Education': PlannedIntervention.Title.EDUCATION,
        'Migration': PlannedIntervention.Title.MIGRATION,
        'Risk Reduction, Climate Adaptation And Recovery': PlannedIntervention.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        'Environmental Sustainability': PlannedIntervention.Title.ENVIRONMENTAL_SUSTAINABILITY,
        'Secretariat Services': PlannedIntervention.Title.SECRETARIAT_SERVICES,
        'National Society Strengthening': PlannedIntervention.Title.NATIONAL_SOCIETY_STRENGTHENING,
        'Multi-purpose Cash': PlannedIntervention.Title.MULTI_PURPOSE_CASH,
        'Community Engagement and Accountability': PlannedIntervention.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,


    }
    return title_dict.get(title)


def parse_disaster_type(disaster_type):
    try:
        return DisasterType.objects.filter(name__icontains=disaster_type).first()
    except ValueError:
        pass


def parse_type_of_onset(onset_type):
    if onset_type == 'Slow':
        return Dref.OnsetType.SLOW
    elif onset_type == 'Sudden':
        return Dref.OnsetType.SUDDEN
    elif onset_type == 'Imminent':
        return Dref.OnsetType.IMMINENT
    return None


def extract_imminent_file(doc, created_by):
    tables = get_table_data(doc)
    paragraphs = get_paragraphs_data(doc)

    # Get item form an array
    def _t(items: List[Any], index=0):
        if items and len(items) > index:
            return items[index]
        return None

    def get_table_cells(table_idx):
        table = tables[table_idx]

        def f(r, c, d=0, as_list=False):
            try:
                if as_list:
                    return table[r][c]
                return _t(table[r][c], d)
            except(IndexError, ValueError):
                pass
        return f
    document = docx.Document(doc)
    data = {}
    # NOTE: Second Paragraph for Country and Region and Dref Title
    paragraph2 = document.paragraphs[1]
    paragraph_element2 = paragraph2._element.xpath('.//w:t')
    data['title'] = get_text_or_null(paragraph_element2)
    if not data['title']:
        raise serializers.ValidationError('A title is required')

    cells = get_table_cells(0)
    data['appeal_code'] = cells(1, 0)
    data['amount_requested'] = parse_string_to_int(cells(1, 1, 1))
    data['disaster_category'] = parse_disaster_category(cells(1, 2))
    data['disaster_type'] = parse_disaster_type(cells(1, 4))
    data['glide_code'] = cells(3, 0)
    data['num_affected'] = parse_string_to_int(cells(3, 1))
    data['num_assisted'] = parse_string_to_int(cells(3, 2))
    data['type_of_onset'] = parse_type_of_onset(cells(5, 0))
    data['date_of_approval'] = parse_date(cells(5, 1))
    data['end_date'] = parse_date(cells(5, 2))
    data['operation_timeframe'] = parse_int(cells(5, 3))
    country = Country.objects.filter(name__icontains=cells(6, 0, 1)).first()
    if country is None:
        raise serializers.ValidationError('A valid country name is required')
    data['country'] = country
    district = cells(6, 2)
    new_district = district.split(',')
    district_list = []
    for d in new_district:
        try:
            district = District.objects.filter(
                country__name__icontains=cells(6, 0, 1),
                name__icontains=d
            ).first()
        except District.DoesNotExist:
            pass
        if district is None:
            continue
        district_list.append(district)
    try:
        if paragraphs[6][0] == 'Approximate date of impact':
            paragraph6 = document.paragraphs[7]._element.xpath('.//w:t')
            event_desc = []
            if len(paragraph6) > 0:
                for desc in paragraph6:
                    event_desc.append(desc.text)
            data['event_text'] = ''.join(event_desc) if event_desc else None
    except IndexError:
        pass
    try:
        if paragraphs[8][0] == 'What is expected to happen?':
            paragraph8 = document.paragraphs[9]._element.xpath('.//w:t')
            event_description = []
            if len(paragraph8) > 0:
                for desc in paragraph8:
                    event_description.append(desc.text)
            data['event_description'] = ''.join(event_description) if event_description else None
    except IndexError:
        pass

    try:
        if paragraphs[11][0] == 'Why your National Society is acting now and what criteria is used to launch this operation.':
            paragraph11 = document.paragraphs[12]._element.xpath('.//w:t')
            anticipatory_actions_desc = []
            if len(paragraph11) > 0:
                for desc in paragraph11:
                    anticipatory_actions_desc.append(desc.text)
        data['anticipatory_actions'] = ''.join(anticipatory_actions_desc) if anticipatory_actions_desc else None
    except IndexError:
        pass
    try:
        if paragraphs[14][0] == 'Scope and scale':
            paragraph13 = document.paragraphs[15]._element.xpath('.//w:t')
            event_scope_description = []
            if len(paragraph13) > 0:
                for desc in paragraph13:
                    event_scope_description.append(desc.text)
            data['event_scope'] = ''.join(event_scope_description) if event_scope_description else None
    except IndexError:
        pass

    # Previous Operation
    # NOTE: I am not sure about the index 1 being used below - Bibek
    table1 = document.tables[1]
    cells = get_table_cells(1)

    data['affect_same_area'] = parse_boolean(cells(0, 1))
    data['affect_same_population'] = parse_boolean(cells(1, 1))
    data['ns_respond'] = parse_boolean(cells(2, 1))
    data['ns_request_fund'] = parse_boolean(cells(3, 1))
    data['ns_request_text'] = cells(4, 1)
    if cells(4, 1) == 'Yes':
        table_one_row_five_column_one = table1.cell(5, 1)._tc.xpath('.//w:t')
        data['dref_recurrent_text'] = get_text_or_null(table_one_row_five_column_one)
    if cells(0, 1) == 'Yes' and \
            cells(1, 1) == 'Yes' and \
            cells(2, 1) == 'Yes' and \
            cells(3, 1) == 'Yes' and \
            cells(4, 1) == 'Yes':
        table_one_row_seven_column_one = table1.cell(7, 1)._tc.xpath('.//w:t')
        recurrent_text = []
        if len(table_one_row_seven_column_one) > 0:
            for desc in table_one_row_seven_column_one:
                recurrent_text.append(desc.text)
        data['dref_recurrent_text'] = ''.join(recurrent_text) if recurrent_text else None

    table_one_row_eight_column_zero = table1.cell(8, 0)._tc.xpath('.//w:t')
    lessons_learned_text = []
    if len(table_one_row_eight_column_zero) > 0:
        for desc in table_one_row_eight_column_zero:
            lessons_learned_text.append(desc.text)
    data['lessons_learned'] = ''.join(lessons_learned_text) if lessons_learned_text else None
    # National Society Actions
    cells = get_table_cells(2)
    # National Society
    national_society_actions = []
    national_society_titles = [
        NationalSocietyAction.Title.NATIONAL_SOCIETY_READINESS,
        NationalSocietyAction.Title.ASSESSMENT,
        NationalSocietyAction.Title.COORDINATION,
        NationalSocietyAction.Title.RESOURCE_MOBILIZATION,
        NationalSocietyAction.Title.ACTIVATION_OF_CONTINGENCY_PLANS,
        NationalSocietyAction.Title.NATIONAL_SOCIETY_EOC,
        NationalSocietyAction.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        NationalSocietyAction.Title.LIVELIHOODS_AND_BASIC_NEEDS,
        NationalSocietyAction.Title.HEALTH,
        NationalSocietyAction.Title.WATER_SANITATION_AND_HYGIENE,
        NationalSocietyAction.Title.PROTECTION_GENDER_AND_INCLUSION,
        NationalSocietyAction.Title.EDUCATION,
        NationalSocietyAction.Title.MIGRATION,
        NationalSocietyAction.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        NationalSocietyAction.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        NationalSocietyAction.Title.ENVIRONMENT_SUSTAINABILITY,
        NationalSocietyAction.Title.MULTI_PURPOSE_CASH,
        NationalSocietyAction.Title.OTHER,
    ]
    for i, title in enumerate(national_society_titles):
        description = cells(i, 1)
        if description:
            national_society_actions.append({
                'title': title,
                'description': description
            })
    # Create national Society objects db level
    national_societies = []
    for national_data in national_society_actions:
        if national_data['description']:
            national = NationalSocietyAction.objects.create(**national_data)
            national_societies.append(national)

    cells = get_table_cells(3)
    ifrc_desc = cells(0, 1, as_list=True) or []

    data['ifrc'] = ''.join(ifrc_desc) if ifrc_desc else None
    partner_national_society_desc = cells(1, 1, as_list=True) or []
    data['partner_national_society'] = ''.join(partner_national_society_desc) if partner_national_society_desc else None

    icrc_desc = cells(2, 1, as_list=True) or []
    data['icrc'] = ''.join(icrc_desc) if icrc_desc else None

    # Other actors
    # Other actors
    cells = get_table_cells(4)
    try:
        if cells(0, 0):
            data['government_requested_assistance'] = parse_boolean(cells(0, 1))
    except (ValueError, IndexError):
        pass
    try:
        national_authorities = cells(1, 2, as_list=True) or []
        data['national_authorities'] = ''.join(national_authorities) if national_authorities else None
    except (ValueError, IndexError):
        pass

    try:
        un_and_other_actors = cells(2, 2, as_list=True) or []
        data['un_or_other_actor'] = ''.join(un_and_other_actors) if un_and_other_actors else None
    except (IndexError, ValueError):
        pass

    try:
        coordination_mechanism = cells(4, 0, as_list=True) or []
        data['major_coordination_mechanism'] = ''.join(coordination_mechanism) if coordination_mechanism else None
    except (IndexError, ValueError):
        pass
    try:
        if paragraphs[28][0] == 'Overall objective of the operation':
            operation_objective = document.paragraphs[29]._element.xpath('.//w:t')
            operation_description = []
            if len(operation_objective) > 0:
                for desc in operation_objective:
                    operation_description.append(desc.text)
            data['operation_objective'] = ''.join(operation_description) if operation_description else None
    except(IndexError, ValueError):
        pass

    # targeting strategy
    try:
        if paragraphs[31][0] == 'Response strategy rationale':
            response_strategy = document.paragraphs[32]._element.xpath('.//w:t')
            response_description = []
            if len(response_strategy) > 0:
                for desc in response_strategy:
                    response_description.append(desc.text)
            data['response_strategy'] = ''.join(response_description) if response_description else None
    except (IndexError, ValueError):
        pass

    try:
        if paragraphs[35][0] == 'Who will be targeted through this operation?':
            paragraph40 = document.paragraphs[36]._element.xpath('.//w:t')
            people_assisted_description = []
            if len(paragraph40) > 0:
                for desc in paragraph40:
                    people_assisted_description.append(desc.text)
            data['people_assisted'] = ''.join(people_assisted_description) if people_assisted_description else None
    except (IndexError, ValueError):
        pass
    try:
        if paragraphs[37][0] == 'Explain the selection criteria for the targeted population':
            paragraph42 = document.paragraphs[38]._element.xpath('.//w:t')
            selection_criteria_description = []
            if len(paragraph42) > 0:
                for desc in paragraph42:
                    selection_criteria_description.append(desc.text)
            data['selection_criteria'] = ''.join(selection_criteria_description) if selection_criteria_description else None
    except (IndexError, ValueError):
        pass

    # NeedsIdentified
    cells = get_table_cells(5)
    needs_identified = []
    needs_titles = [
        IdentifiedNeed.Title.SHELTER_HOUSING_AND_SETTLEMENTS,
        IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        IdentifiedNeed.Title.HEALTH,
        IdentifiedNeed.Title.WATER_SANITATION_AND_HYGIENE,
        IdentifiedNeed.Title.PROTECTION_GENDER_AND_INCLUSION,
        IdentifiedNeed.Title.EDUCATION,
        IdentifiedNeed.Title.RISK_REDUCTION_CLIMATE_ADAPTATION_AND_RECOVERY,
        IdentifiedNeed.Title.COMMUNITY_ENGAGEMENT_AND_ACCOUNTABILITY,
        IdentifiedNeed.Title.ENVIRONMENT_SUSTAINABILITY,
        IdentifiedNeed.Title.SHELTER_CLUSTER_COORDINATION,
    ]
    for i, needs_title in enumerate(needs_titles):
        try:
            description = cells(i * 2 + 1, 0)  # Use Every alternate row, so i*2 + 1
            if description:
                data_new = {
                    'title': needs_title,
                    'description': description
                }
                needs_identified.append(data_new)
        except IndexError:
            pass

    needs = []
    for need in needs_identified:
        if need['description']:
            identified = IdentifiedNeed.objects.create(**need)
            needs.append(identified)

    # Targeting Population
    cells = get_table_cells(6)
    data['women'] = parse_string_to_int(cells(0, 1))
    data['girls'] = parse_string_to_int(cells(1, 1))
    data['men'] = parse_string_to_int(cells(2, 1))
    data['boys'] = parse_string_to_int(cells(3, 1))
    data['total_targeted_population'] = parse_string_to_int(cells(4, 1))
    data['people_per_local'] = parse_float(cells(1, 2))
    data['people_per_urban'] = parse_float(cells(1, 3))
    data['disability_people_per'] = parse_float(cells(3, 2))
    data['people_targeted_with_early_actions'] = parse_int(cells(4, 4))

    # Risk And Security Considerations
    cells = get_table_cells(7)
    mitigation_actions = []
    for i in range(2, 5):
        risk = cells(i, 0)
        mitigation = cells(i, 1)
        data_row = {
            'risk': risk,
            'mitigation': mitigation,
        }
        mitigation_actions.append(data_row)
    mitigation_list = []
    for action in mitigation_actions:
        if action['risk'] or action['mitigation']:
            risk_security = RiskSecurity.objects.create(**action)
            mitigation_list.append(risk_security)
    data['risk_security_concern'] = cells(6, 0)

    # PlannedIntervention Table
    planned_intervention = []
    for i in range(8, 20):
        cells = get_table_cells(i)
        title = cells(0, 1, 0)
        budget = cells(0, 3, 1)
        targeted_population = cells(1, 3)

        indicators = []
        for j in range(3, 6):
            indicators.append({
                'title': cells(j, 0),
                'target': parse_int(cells(j, 2)),
            })
        indicators_object_list = []
        for indicator in indicators:
            if indicator['title'] is None:
                continue
            planned_object = PlannedInterventionIndicators.objects.create(**indicator)
            indicators_object_list.append(planned_object)

        priority_description = cells(6, 1)
        planned_data = {
            'title': parse_planned_intervention_title(title),
            'budget': parse_int(budget),
            'person_targeted': parse_int(targeted_population),
            'description': priority_description
        }
        if planned_data['budget'] or planned_data['person_targeted'] or planned_data['description']:
            planned = PlannedIntervention.objects.create(**planned_data)
            planned.indicators.add(*indicators_object_list)
            planned_intervention.append(planned)

    # About Support Service
    try:
        if paragraphs[61][0] == 'How many volunteers and staff involved in the response? Briefly describe their role.':
            paragraph53 = document.paragraphs[62]._element.xpath('.//w:t')
            human_resource_description = []

            if len(paragraph53) > 0:
                for desc in paragraph53:
                    human_resource_description.append(desc.text)
            data['human_resource'] = ''.join(human_resource_description) if human_resource_description else None
    except (IndexError, ValueError):
        pass

    try:
        if paragraphs[63][0] == 'Will surge personnel be deployed? Please provide the role profile needed.':
            paragraph55 = document.paragraphs[64]._element.xpath('.//w:t')
            surge_personnel_deployed_description = []
            if len(paragraph55) > 0:
                for desc in paragraph55:
                    surge_personnel_deployed_description.append(desc.text)
            data['surge_personnel_deployed'] = ''.join(surge_personnel_deployed_description) if surge_personnel_deployed_description else None
    except (IndexError, ValueError):
        pass
    try:
        if paragraphs[65][0] == 'If there is procurement, will it be done by National Society or IFRC?':
            paragraph70 = document.paragraphs[66]._element.xpath('.//w:t')
            logistic_capacity_of_ns_description = []
            if len(paragraph70) > 0:
                for desc in paragraph70:
                    logistic_capacity_of_ns_description.append(desc.text)
            data['logistic_capacity_of_ns'] = ''.join(logistic_capacity_of_ns_description) if logistic_capacity_of_ns_description else None
    except (IndexError, ValueError):
        pass

    try:
        if paragraphs[68] == 'How will this operation be monitored?':
            paragraph73 = document.paragraphs[69]._element.xpath('.//w:t')
            pmer_description = []
            if len(paragraph73) > 0:
                for desc in paragraph73:
                    pmer_description.append(desc.text)
            data['pmer'] = ''.join(pmer_description) if pmer_description else None
    except(IndexError, ValueError):
        pass

    try:
        if paragraphs[70][0] == 'Please briefly explain the National Societies communication strategy for this operation.':
            paragraph75 = document.paragraphs[71]._element.xpath('.//w:t')
            communication_description = []
            if len(paragraph75) > 0:
                for desc in paragraph75:
                    communication_description.append(desc.text)
            data['communication'] = ''.join(communication_description) if communication_description else None
    except(IndexError, ValueError):
        pass

    # Contact Information
    try:
        national_society_contact = parse_contact_information(paragraphs[76] or [])
        data['national_society_contact_title'] = national_society_contact[2] if national_society_contact[2] and national_society_contact[1] != ", , " else None
        data['national_society_contact_email'] = national_society_contact[4] if national_society_contact[4] else None
        data['national_society_contact_phone_number'] = national_society_contact[6] if national_society_contact[6] else None
        data['national_society_contact_name'] = national_society_contact[0] if national_society_contact[0] else None
    except IndexError:
        pass

    try:

        ifrc_appeal_manager = parse_contact_information(paragraphs[77] or [])
        data['ifrc_appeal_manager_title'] = ifrc_appeal_manager[2] if ifrc_appeal_manager[2] else None
        data['ifrc_appeal_manager_email'] = ifrc_appeal_manager[4] if ifrc_appeal_manager[4] else None
        data['ifrc_appeal_manager_phone_number'] = ifrc_appeal_manager[6] if ifrc_appeal_manager[6] else None
        data['ifrc_appeal_manager_name'] = ifrc_appeal_manager[0] if ifrc_appeal_manager[0] else None
    except IndexError:
        pass

    try:
        ifrc_project_manager = parse_contact_information(paragraphs[78] or [])
        data['ifrc_project_manager_title'] = ifrc_project_manager[2] if ifrc_project_manager[2] else None
        data['ifrc_project_manager_email'] = ifrc_project_manager[4] if ifrc_project_manager[4] else None
        data['ifrc_project_manager_phone_number'] = ifrc_project_manager[6] if ifrc_project_manager[6] else None
        data['ifrc_project_manager_name'] = ifrc_project_manager[0] if ifrc_project_manager[0] else None
    except IndexError:
        pass

    try:
        ifrc_emergency = parse_contact_information(paragraphs[79] or [])
        data['ifrc_emergency_title'] = ifrc_emergency[2] if ifrc_emergency[2] else None
        data['ifrc_emergency_email'] = ifrc_emergency[4] if ifrc_emergency[4] else None
        data['ifrc_emergency_phone_number'] = ifrc_emergency[6] if ifrc_emergency[6] else None
        data['ifrc_emergency_name'] = ifrc_emergency[0] if ifrc_emergency[0] else None
    except IndexError:
        pass

    try:
        media = parse_contact_information(paragraphs[80] or [])
        data['media_contact_title'] = media[2] if media[2] else None
        data['media_contact_email'] = media[4] if media[4] else None
        data['media_contact_phone_number'] = media[6] if media[6] else None
        data['media_contact_name'] = media[0] if media[0] else None
    except IndexError:
        pass

    data['is_published'] = False
    data['national_society'] = country
    data['created_by'] = created_by

    dref = Dref.objects.create(**data)
    dref.planned_interventions.add(*planned_intervention)
    dref.needs_identified.add(*needs)
    dref.national_society_actions.add(*national_societies)
    dref.risk_security.add(*mitigation_list)
    if len(district_list) > 0 and None not in district_list:
        dref.district.add(*district_list)
    return dref
